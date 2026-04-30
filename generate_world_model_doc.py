from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'PingFang SC'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'PingFang SC')

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hf = hs.font
    hf.name = 'PingFang SC'
    hf.bold = True
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'PingFang SC')

sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

def add_title(text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(26)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True

def add_h1(text):
    doc.add_heading(text, level=1)

def add_h2(text):
    doc.add_heading(text, level=2)

def add_h3(text):
    doc.add_heading(text, level=3)

def add_para(text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_highlight(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0x0d, 0x47, 0xa1)
    run.font.size = Pt(11)
    return p

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Menlo'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pf = p.paragraph_format
    pf.left_indent = Cm(1)
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    return table

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

# ============================================================
# 封面
# ============================================================
add_title('王牌顾问 · 世界模型理论深度分析')
add_subtitle('—— 手工世界模型与前沿 AI 理论的映射、对话与反思')
doc.add_paragraph()
doc.add_paragraph()

# ============================================================
# 引言
# ============================================================
add_h1('引言：为什么一个游戏值得用世界模型理论来审视')

add_para('"王牌顾问"是一个用代码还原上海二手房经营的模拟游戏——21 天、6 套房源、20+ 个引擎、24 步流水线、120+ 个平衡常量。从纯工程角度看，它是一个精心设计的 DDD 系统；从游戏设计角度看，它是一个数值驱动的经营模拟。但如果戴上 AI 研究者的眼镜，我们会看到一个更有趣的东西：')
add_highlight('它是一个完全由人类工程师手工构建的"世界模型"（World Model）——一个关于二手房市场的内部表征，能够生成未来状态的预测，并据此做出评价。')
add_para('2018 年，Ha 和 Schmidhuber 发表了划时代的论文《World Models》，提出了 VAE + MDN-RNN + Controller 的三组件架构。2022 年起，LeCun 提出了 JEPA（Joint-Embedding Predictive Architecture）框架，主张在抽象表征空间而非像素空间做预测。Dreamer 系列则证明了"在世界模型的想象中学习"的可行性。')
add_para('这些理论与我们的代码有什么关系？它们不仅是学术概念——它们提供了一套精确的语言，帮助我们理解自己在做什么、做得怎么样、以及还可以做什么。')
add_para('本文分 9 个章节，从理论映射、结构分析、保真度评估到设计启示，试图回答一个核心问题：')

add_highlight('一个手工构建的世界模型，与一个通过学习获得的世界模型，它们之间的距离有多远？这个距离告诉我们什么？')

# ============================================================
# 第一章
# ============================================================
add_h1('第一章 两条路线：世界模型理论的两个范式')

add_h2('1.1 Ha-Schmidhuber 路线：生成式世界模型')

add_para('2018 年，David Ha 和 Jürgen Schmidhuber 在《World Models》中提出了一个优雅的架构：')
add_code('        环境（真实世界或游戏）')
add_code('              │')
add_code('        V（VAE）── 观测压缩 → 潜在向量 z')
add_code('              │')
add_code('        M（MDN-RNN）── 序列预测 → 下一个 z 的分布')
add_code('              │')
add_code('        C（Controller）── 根据 z 做决策')
add_para('核心思想是：智能体不需要直接在真实环境中试错，而是可以在一个学到的"梦境"（dream）中训练策略。VAE 负责把高维观测（如游戏画面）压缩成低维潜在表征，MDN-RNN 负责预测潜在空间中的状态转移，Controller 负责在预测的基础上做决策。')
add_para('这条路的关键假设是：', bold_prefix='')
add_bullet('观测空间是高维的（像素级），需要压缩')
add_bullet('转移动力学可以用序列模型近似')
add_bullet('在"梦境"中训练的策略可以迁移到真实环境')

add_h2('1.2 LeCun JEPA 路线：表征预测式世界模型')

add_para('2022 年起，Yann LeCun 在多篇论文和演讲中提出了一个不同的框架——JEPA（Joint-Embedding Predictive Architecture）。其核心主张是：')
add_highlight('不要预测像素，要预测表征。')
add_para('JEPA 的架构是：')
add_code('        环境')
add_code('          │')
add_code('    S（State Encoder）── 编码当前状态 → 表征 s')
add_code('          │')
add_code('    P（Predictor）── 在表征空间中预测 → s\' = P(s, a)')
add_code('          │')
add_code('    C（Cost Module）── 评估预测状态的"好坏" → cost')
add_para('JEPA 与 Ha-Schmidhuber 的关键区别在于：')
add_table(
    ['维度', 'Ha-Schmidhuber', 'LeCun JEPA'],
    [
        ['预测什么', '像素/观测的分布', '抽象表征的向量'],
        ['在哪里预测', '观测空间（高维）', '表征空间（低维抽象）'],
        ['如何处理不确定性', '混合密度网络（概率分布）', '不建模不确定性（点估计）'],
        ['代价函数', '重建误差', '抽象表征的预测误差'],
        ['哲学立场', '世界是可生成的', '世界是可抽象的'],
    ]
)

add_h2('1.3 Dreamer 路线：在想象中强化学习')

add_para('Dreamer（Hafner et al., 2020-2023）是 Ha-Schmidhuber 路线的工程化巅峰。DreamerV3 的架构是：')
add_code('        环境交互 → 真实轨迹')
add_code('          │')
add_code('    RSSM（循环状态空间模型）── 编码/预测状态')
add_code('          │')
add_code('    世界模型（Dreamer 学到的）── 在潜在空间中"想象"轨迹')
add_code('          │')
add_code('    Actor-Critic ── 在想象轨迹上训练策略')
add_para('Dreamer 的核心突破是证明了：一个足够好的世界模型 + 在其想象中训练的策略，可以在 200M 帧内解决 Atari、DMC 等基准任务，且效率远高于 model-free 方法。')

add_h2('1.4 三条路线的共同前提')
add_para('无论 Ha-Schmidhuber、JEPA 还是 Dreamer，它们共享一个前提：')
add_highlight('世界模型是对环境的内部表征，能够生成（或预测）未来状态，并用于决策或评价。')
add_para('这个前提与王牌顾问的关系是什么？让我们逐一映射。')

# ============================================================
# 第二章
# ============================================================
add_h1('第二章 V/M/C 映射：我们的代码结构对应什么')

add_h2('2.1 状态编码器 V——GameState 就是我们的"表征空间"')
add_para('在 Ha-Schmidhuber 中，V（VAE）负责把高维观测压缩为低维潜在向量 z。在王牌顾问中，我们没有像素级的观测需要压缩——我们的"观测"从一开始就是结构化的数值。')
add_para('换句话说，', bold_prefix='')
add_highlight('王牌顾问的世界模型跳过了 V（编码器），直接在人类设计的"表征空间"中工作。GameState 本身就是 z。')
add_para('但这并不意味着没有编码工作。每一个 GameState 字段都是一个设计决策——选择保留什么、丢弃什么、如何量化。以 CaseItem 为例：')
add_table(
    ['保留的维度', '量化方式', '丢弃的信息'],
    [
        ['trust（信任）', '0-100 标量', '业主的性格细节、沟通语气、历史情绪波动'],
        ['heat（热度）', '0-100 标量', '客户的真实看房频率、带看质量、看房时长'],
        ['patience（耐心）', '0-100 标量', '业主的生活背景、卖房动机的复杂度'],
        ['windowDays（窗口期）', '整数天', '业主的隐性心理底线、家人意见分歧'],
        ['competitiveness（竞争力）', '0-100 综合值', '竞品房源的故事细节、门店的公关能力'],
        ['d1/d2/d3（三轴）', '各 0-100', '客户的隐性偏好、业主的隐性底线'],
    ]
)
add_para('这种编码决策与 VAE 的信息瓶颈（information bottleneck）原理相同：在有限的维度中保留最大化的决策相关信息，同时丢弃噪声。区别在于 VAE 通过训练自动学习瓶颈，而我们通过领域知识手工设计瓶颈。')

add_h2('2.2 转移模型 M——24 步流水线就是我们的 MDN-RNN')
add_para('在 Ha-Schmidhuber 中，M（MDN-RNN）负责预测下一个潜在状态 z\' 的概率分布。在王牌顾问中，这个角色由 resolveOneDay 的 24 步流水线承担：')
add_code('z_t = GameState(day=t)')
add_code('z_{t+1} = resolveOneDay(z_t, actions_t, rngState)')
add_para('但这里有一个根本性的差异：MDN-RNN 输出的是一个概率分布（混合密度网络），而 resolveOneDay 输出的是一个确定性状态（加上 randomInt 的微小扰动）。')
add_para('我们可以将 resolveOneDay 的 24 步拆解为 M 的子组件：')
add_table(
    ['流水线阶段', 'M 的子功能', '理论对应'],
    [
        ['步骤 1-3：市场脉冲', '环境动力学', '自然状态转移（不需要动作）'],
        ['步骤 4-7：竞争注入', '对手模型', '多智能体环境中的 others\' policy'],
        ['步骤 8-13：客户反应', '响应模型', '条件转移 P(s\'|s,a,others)'],
        ['步骤 14-16：机会/竞争结算', '终端状态判定', '吸收态（absorbing state）检测'],
        ['步骤 17-19：成交/抢房', '奖励触发', '稀疏奖励的发放点'],
        ['步骤 20-24：衰减/生成/复盘', '状态重置', 'episode 内的时间步推进'],
    ]
)

add_h2('2.3 控制器 C——玩家就是我们的 Controller')
add_para('在 Ha-Schmidhuber 中，Controller 是一个神经网络，接收 z 和 M 的隐状态，输出动作。在王牌顾问中，控制器是人类玩家——但游戏系统通过 Action 机制对玩家施加了结构化约束：')
add_code('Action = {')
add_code('  actionType: "owner-maintain" | "show-property" | "price-advice" | ...')
add_code('  costEnergy: number       // 精力消耗')
add_code('  costPromotionBudget: number  // 推广金消耗')
add_code('  targetCaseId: string     // 动作目标')
add_code('}')
add_para('每个 Action 通过 actionResolvers.ts 中的 30+ 个 resolver 映射到 GameState 的状态变更。这相当于一个结构化的动作空间（structured action space）——玩家不能输出任意的向量，只能从预定义的离散动作集中选择。')
add_highlight('映射总结：V（编码器）= 领域模型中的类型定义和量化决策；M（转移模型）= resolveOneDay 的 24 步流水线；C（控制器）= 人类玩家 + Action 机制。三者共同构成一个完整的世界模型回路。')

# ============================================================
# 第三章
# ============================================================
add_h1('第三章 JEPA 视角：抽象层次与丢弃策略')

add_h2('3.1 JEPA 的核心洞见：预测表征而非像素')
add_para('LeCun 的 JEPA 框架最深刻的洞见是：智能体应该在抽象表征空间中做预测，而不是在观测空间中做预测。这意味着世界模型的关键不在于"重建输入"，而在于"提取与决策相关的信息"。')
add_para('王牌顾问的领域模型天然就是一个 JEPA 风格的表征系统——它从不试图还原真实二手房交易的每一个细节，而是提取与经纪决策相关的抽象维度。')

add_h2('3.2 四层抽象阶梯')
add_para('从真实世界到游戏内的最终评分，信息经历了四层抽象：')
add_code('        真实世界（无限维）')
add_code('              │  第一层：领域建模（丢弃 99.9% 的信息）')
add_code('        结构化数据（Case/Customer/Market 30+ 字段）')
add_code('              │  第二层：评分压缩（D1/D2/D3 三轴）')
add_code('        竞争力评估（3 个 0-100 值）')
add_code('              │  第三层：终局评判（ability/defense/satisfaction）')
add_code('        最终评分（100 分制）')
add_code('              │  第四层：等级映射（王牌/漂亮/过线/没保住）')
add_code('        四级评价（离散标签）')
add_para('每一层都是一次信息压缩，每一次压缩都是一次"丢弃什么"的决策。JEPA 的理论提醒我们：好的世界模型不是保留最多信息的模型，而是保留与目标最相关信息的模型。')

add_h2('3.3 逐层分析：我们丢弃了什么')

add_h3('第一层：从真实世界到结构化数据')
add_para('这一层由 models.ts 中的类型定义完成。丢弃的信息包括：')
add_bullet('业主的人格故事（为什么卖房、家庭情况、情感状态）')
add_bullet('房源的空间感受（采光的主观体验、邻居的真实关系）')
add_bullet('市场的微观结构（具体每个买家的看房历史、每笔交易的谈判细节）')
add_bullet('经纪人的专业直觉（"这套房适合谁"的隐性知识）')
add_para('保留的信息包括：')
add_bullet('信任、热度、耐心等可量化的互动指标')
add_bullet('匹配度、购买意向、置信度等可计算的行为指标')
add_bullet('需求热度、供给压力等可聚合的市场指标')
add_highlight('设计启示：这一层的抽象质量决定了整个世界模型的上限。如果丢弃的信息中包含决策关键因素（如业主的真实底线），那么无论下游引擎多么精巧，都无法弥补这个信息损失。')

add_h3('第二层：从结构化数据到三轴竞争力')
add_para('这一层由 scoring.ts 完成。核心公式是：')
add_code('competitiveness = D1 × 0.5 + D2 × 0.25 + D3 × 0.25')
add_para('D1 使用了 log2 平滑，D2 是 7 轴加权，D3 是多信号线性组合。关键的丢弃决策包括：')
add_bullet('D1 丢弃了单个客户的具体属性，只保留漏斗的统计特征（厚度、速度、停滞率）')
add_bullet('D2 丢弃了房源的动态变化（因为禀赋是固定的），只保留静态品质')
add_bullet('D3 丢弃了业主的精细行为模式（如"只在周三打电话才接"），只保留宏观配合度')
add_para('这个压缩比是 30:3——从 30+ 个字段压缩到 3 个值。这是一个激进的压缩，但它保留了三个最具决策价值的维度："有没有人在看"、"房子本身好不好"、"业主愿不愿意配合"。')

add_h3('第三层：从三轴竞争力到终局评判')
add_para('这一层由 resultEvaluation.ts 的三轴终评完成。关键的丢弃决策包括：')
add_bullet('Ability 丢弃了"卖了多少钱"的连续值，压缩为 outrun/flat/lose 三档')
add_bullet('Defense 丢弃了"差多少被抢"的连续值，压缩为 held/at_risk/lost 三档')
add_bullet('Satisfaction 丢弃了"信任是多少"的连续值，压缩为 5 档满意度')

add_h2('3.4 JEPA 的"联合嵌入"在我们系统中的对应')
add_para('JEPA 的"联合嵌入"（Joint-Embedding）指的是：编码器同时编码当前状态和未来状态，使得两个表征在同一空间中可比较。')
add_para('在王牌顾问中，这个对应关系非常微妙但确实存在：')
add_bullet('competitivenessSnapshots 记录了每天的竞争力快照——这是状态的时间序列编码')
add_bullet('competitivenessDelta（今天的竞争力 - 昨天的竞争力）用于 urgency 计算——这是表征空间中的差分预测')
add_bullet('storylineState（healthy/fragile/sliding/critical）是基于连续状态的离散聚类——这是表征空间中的区域划分')
add_highlight('我们的系统没有明确训练一个联合嵌入编码器，但 competitivenessSnapshots + competitivenessDelta 的设计，本质上实现了 JEPA 的"在表征空间中比较不同时刻"的思想。')

# ============================================================
# 第四章
# ============================================================
add_h1('第四章 手工世界模型 vs 学习型世界模型')

add_h2('4.1 两种构建路径的本质差异')
add_para('王牌顾问的世界模型是完全手工构建的——每一个公式、每一个常量、每一个阈值都是工程师的领域知识的编码。这与 Ha-Schmidhuber、Dreamer 等学习型世界模型形成了鲜明对比：')
add_table(
    ['维度', '手工世界模型（王牌顾问）', '学习型世界模型（Dreamer）'],
    [
        ['构建方式', '工程师编写公式和规则', '神经网络从数据中学习'],
        ['知识来源', '领域专家的直觉和经验', '环境交互的观测数据'],
        ['状态空间', '人工设计的结构化类型', '自动学到的潜在向量'],
        ['转移函数', '确定性公式 + 小幅随机扰动', '概率分布（MDN/扩散模型）'],
        ['可解释性', '极高（每个常量都有业务含义）', '极低（黑盒潜在空间）'],
        ['修改成本', '改一行公式即可', '需要重新训练'],
        ['泛化能力', '限于设计者预见的场景', '理论上可泛化到训练分布内'],
        ['精度上限', '受限于设计者的领域知识', '受限于训练数据的质量和数量'],
    ]
)

add_h2('4.2 手工世界模型的优势')
add_para('王牌顾问的 120+ 个平衡常量不是劣势，而是精心编码的领域知识。每个常量都承载着一个业务洞察：')
add_bullet('trustDecayMultiplier 的 55% 跨度（焦虑型 1.4 vs 等价型 0.9）编码了"不同业主对忽视的敏感度差异"', bold_prefix='洞察 1：')
add_bullet('marketDealSlots 机制编码了"二手房市场是零和博弈"的业务真相', bold_prefix='洞察 2：')
add_bullet('D1 的 log2 边际递减编码了"客户数量的价值不线性增长"的业务常识', bold_prefix='洞察 3：')
add_bullet('rivalLossProbabilityScale 从 0.08 到 0.9 编码了"难度梯度应该影响被撬概率"的设计意图', bold_prefix='洞察 4：')
add_para('如果用学习型世界模型来替代，这些洞察需要从大量的经纪人行为数据中学习——数据获取成本极高，且学到的表征可能不如手工设计的可解释。')

add_h2('4.3 手工世界模型的局限')
add_para('然而，手工世界模型也面临三个根本性限制：')

add_h3('限制一：设计者的认知边界')
add_para('所有公式都是设计者对"二手房市场如何运作"的建模假设。如果设计者的理解有偏差，模型就会有系统性偏差。例如：')
add_bullet('当前 D1 只看客户数量，不看客户质量——一个高质量客户（预算充足、决策果断）和一个低质量客户（犹豫不决、预算不足）在 D1 中贡献相同')
add_bullet('当前信任衰减是线性的，没有建模"信任崩溃是突然的"这个心理现实')

add_h3('限制二：组合爆炸的应对能力')
add_para('手工公式擅长建模单变量关系（如"超价 5% 则信任下降"），但在多变量交互时容易变得复杂难维护。例如：')
add_code('当业主是焦虑型（trustDecayMultiplier=1.4）AND')
add_code('  报价超市场价 10%（priceElasticity=1.25）AND')
add_code('  竞品压力 > 34（rivalPressure）AND')
add_code('  窗口期 < 3 天（windowDays）')
add_code('→ 四变量交互下的衰减速度是多少？')
add_para('手工公式只能通过乘法/加法线性组合来近似，而学习型模型可以自动发现非线性交互。')

add_h3('限制三：缺乏反事实推理能力')
add_para('Dreamer 的核心优势是能在世界模型中"想象"不同行动的后果（counterfactual simulation）。王牌顾问的世界模型虽然也支持"如果我今天维护业主，明天 trust 会怎样"的推理，但这种推理是：')
add_bullet('前向的（只能从当前状态向前推演）')
add_bullet('单路径的（只推演了一条随机种子下的轨迹）')
add_bullet('非优化的（没有自动搜索最优动作序列）')
add_highlight('手工世界模型的优势在于精确性和可解释性，学习型世界模型的优势在于泛化能力和反事实推理。两者不是替代关系，而是互补关系。')

# ============================================================
# 第五章
# ============================================================
add_h1('第五章 Dreamer 的"在想象中学习"——适用性分析')

add_h2('5.1 Dreamer 的核心机制')
add_para('Dreamer 的学习流程是：')
add_code('1. 与真实环境交互，收集真实轨迹')
add_code('2. 用真实轨迹训练世界模型（RSSM）')
add_code('3. 在世界模型的"想象"中生成虚拟轨迹')
add_code('4. 在虚拟轨迹上训练 Actor-Critic 策略')
add_code('5. 用学到的策略与真实环境交互 → 回到步骤 1')
add_para('关键问题是：如果用王牌顾问的真实环境替换 Dreamer 的 Atari/DMC 环境，这个流程能工作吗？')

add_h2('5.2 理论可行性分析')
add_para('答案是：理论上可以，但有几个关键前提需要满足。')

add_h3('前提一：可微性')
add_para('Dreamer 的世界模型训练需要反向传播梯度。王牌顾问的 resolveOneDay 包含 randomInt、chance 等随机操作和 clamp、max/min 等非光滑操作。这些操作在技术上可以通过重参数化技巧（reparameterization trick）和软化近似（soft approximation）来处理，但工程成本不低。')

add_h3('前提二：状态空间的维度')
add_para('王牌顾问的 GameState 有 100+ 个标量字段，加上动态长度的数组（cases[]、opportunities[]、customers[]）。这个状态空间比 Atari 的像素空间小得多，但结构更复杂（混合了标量、枚举、嵌套对象）。需要一个能处理异构状态的编码器。')

add_h3('前提三：奖励信号的稀疏性')
add_para('王牌顾问的终局奖励（100 分制评分）是极度稀疏的——21 天只在最后一天给一个奖励。Dreamer 的 Actor-Critic 在稀疏奖励下需要更长的想象轨迹才能有效传播价值信号。')
add_para('但系统中存在中间信号可以作为 dense reward 使用：')
add_bullet('每日竞争力变化（competitivenessDelta）')
add_bullet('客户阶段升格事件')
add_bullet('成交事件')
add_bullet('信任/热度的每日变化')

add_h2('5.3 如果真的用 Dreamer 学习，会怎样？')
add_para('假设我们用 DreamerV3 来学习王牌顾问的策略，以下是一些可预见的结果：')
add_table(
    ['方面', 'Dreamer 预期表现', '人类玩家预期表现'],
    [
        ['市场时机把握', '可能更好（能学到正弦波周期）', '依赖直觉，可能错过周期'],
        ['多房源资源分配', '可能更好（全局优化）', '受注意力限制，容易偏科'],
        ['业主关系维护', '可能更机械（没有情感理解）', '能根据对话调整策略'],
        ['应变能力', '可能更差（没见过的分布外场景）', '能灵活应对意外'],
        ['可解释性', '极低（黑盒策略）', '可以解释每个决策的理由'],
    ]
)
add_highlight('Dreamer 在王牌顾问上最可能的优势是：全局资源分配和市场时机把握——这两个恰好是人类最弱的地方。最可能的劣势是：应变能力和可解释性——这两个恰好是人类最强的地方。')

# ============================================================
# 第六章
# ============================================================
add_h1('第六章 Cost Module 与评分体系的对齐')

add_h2('6.1 LeCun 的 Cost Module 理论')
add_para('在 LeCun 的 JEPA 框架中，Cost Module 是整个系统的评价核心。它接收当前状态的表征 s 和预测的未来状态表征 s\'，输出一个标量 cost，表示"这个未来状态有多不好"。')
add_para('LeCun 把 Cost Module 分为两类：')
add_bullet('内在成本（intrinsic cost）：与当前任务直接相关的即时代价，如"离目标的距离"', bold_prefix='')
add_bullet('可训练成本（trainable cost）：通过学习获得的长期评价，如"这个状态对长远目标的影响"', bold_prefix='')

add_h2('6.2 王牌顾问的 Cost Module 解剖')
add_para('王牌顾问的"Cost Module"是 resultEvaluation.ts 中的三轴评分系统。我们可以将其精确映射到 LeCun 的框架：')
add_table(
    ['LeCun 概念', '王牌顾问对应', '代码实现'],
    [
        ['内在成本（intrinsic cost）', '每日竞争力衰减', 'tickCases 中的 trust/heat/patience 衰减'],
        ['可训练成本（trainable cost）', '终局三轴评分', 'ability/defense/satisfaction'],
        ['目标条件（goal condition）', 'difficultyTargets', 'targetScore + goalTier 权重'],
        ['动作代价（action cost）', 'costEnergy + costPromotionBudget', 'Action 接口的资源消耗'],
        ['约束违反（constraint violation）', '不变量检查', 'collectInvariantAlerts'],
    ]
)

add_h2('6.3 三轴评分的理论解读')
add_para('三轴评分（ability 40 + defense 35 + satisfaction 25 = 100）可以被理解为一个多目标 Cost Module，每个轴对应一个独立的优化目标：')

add_h3('Ability（能力分，满分 40）——"卖得好不好"')
add_para('这是 LeCun 框架中与任务最直接相关的 cost 维度。它的计算基于：')
add_code('base = outrun ? 1 : flat ? 0.65 : 0.2')
add_para('outrun（卖出价 ≥ 市场价 × 98.5%）对应"高质量完成任务"，flat（≥ 95%）对应"基本完成"，lose 对应"完成但代价大"。')
add_para('从世界模型理论角度看，Ability 是对"转移轨迹终端质量"的评价——它不关心过程，只关心最终的卖出价格。')

add_h3('Defense（防守分，满分 35）——"有没有被抢"')
add_para('这是一个"避免负奖励"的目标。held（窗口期 > 3 天且信任 > 50）对应"安全"，at_risk（窗口期 ≤ 3 或信任 ≤ 50）对应"危险"，lost 对应"失败"。')
add_para('从世界模型理论角度看，Defense 是对"到达吸收态（absorbing state）的风险"的评价——房源被竞品抢走是一个不可逆的负向吸收态。')

add_h3('Satisfaction（满意度，满分 25）——"业主开不开心"')
add_para('这是一个"软约束"目标——即使卖掉了房子，如果业主不开心，分数也会降低。happy/neutral/no_regret/regret/unhappy 五档构成了一个渐进的满意度梯度。')
add_para('从世界模型理论角度看，Satisfaction 是对"状态转移过程中信任关系维护质量"的评价——它关注的不是结果，而是过程中人与人之间的信任状态。')

add_h2('6.4 goalTier 权重的理论含义')
add_para('goalTier 权重（core=1.0 / important=0.7 / normal=0.4）是 Cost Module 的条件化（conditioning）机制。它等价于 LeCun 框架中的"目标条件"（goal condition）——不同的房源有不同的评价权重，这意味着"把 anchor 房卖好"比"把 sacrifice 房卖好"的 cost 降低更多。')
add_highlight('从理论角度看，goalTier 是一种多任务学习（multi-task learning）的简化形式——同一局中的不同房源可以被视为不同的子任务，权重反映了子任务的重要性差异。')

# ============================================================
# 第七章
# ============================================================
add_h1('第七章 保真度评估：我们的世界模型有多"真"')

add_h2('7.1 保真度的三个维度')
add_para('世界模型的保真度（fidelity）可以从三个维度评估：')
add_bullet('结构保真度：模型的组件结构是否反映了真实世界的结构', bold_prefix='')
add_bullet('动力学保真度：模型的状态转移是否反映了真实世界的变化规律', bold_prefix='')
add_bullet('评价保真度：模型的 cost function 是否反映了真实世界的"好/坏"判断', bold_prefix='')

add_h2('7.2 结构保真度：约 70%')
add_para('王牌顾问的结构设计高度还原了二手房经营的核心组件：')
add_table(
    ['真实世界组件', '游戏模型对应', '保真度', '缺失'],
    [
        ['业主', 'Case + OwnerArchetype', '高', '多人决策（夫妻/父母意见）'],
        ['客户', 'CustomerProfile + RuntimeState', '高', '隐性偏好（学区焦虑等）'],
        ['市场', 'MarketCell × 2', '中', '更多商圈维度、政策冲击'],
        ['竞品', 'RivalStore + RivalListing', '中', '竞品的人际网络和信息战'],
        ['公司', 'CompanyPressure', '低', '团队协作、领导关系、培训体系'],
        ['经纪人', 'Player（隐式）', '低', '个人成长、情绪管理、专业积累'],
    ]
)
add_para('结构保真度的核心缺失是"人"的维度——业主和客户被简化为数值，丢失了人格、情感、社会关系等无法量化的因素。这不是疏忽，而是有意的抽象——但抽象的边界决定了模型的表达力上限。')

add_h2('7.3 动力学保真度：约 60%')
add_para('动力学保真度评估的是"状态转移规律是否符合现实"。逐项分析：')

add_h3('高度保真的动力学')
add_bullet('信任衰减的差异化（trustDecayMultiplier × 人格）——现实中确实有"一碰就碎"和"怎么都行"的业主', bold_prefix='')
add_bullet('客户阶段的渐进性（7 阶段漏斗）——现实中买房确实是一个从"了解"到"出价"的渐进过程', bold_prefix='')
add_bullet('市场周期的正弦波——上海二手房市场确实有季节性波动', bold_prefix='')

add_h3('中度保真的动力学')
add_bullet('竞品压力的阈值效应（pressure ≥ 34 才触发）——现实中竞品影响确实是渐进的，但阈值效应过于机械化', bold_prefix='')
add_bullet('成交概率的两日结算——现实中确实有"今天看中，明天再谈"的模式，但两天的时序过于简化', bold_prefix='')

add_h3('低度保真的动力学')
add_bullet('信任衰减的线性假设——现实中信任的崩塌往往是非线性的（"最后一根稻草"效应）', bold_prefix='')
add_bullet('客户流失的马尔可夫性——现实中客户的流失决策受到历史记忆、社会压力等非马尔可夫因素影响', bold_prefix='')
add_bullet('成交名额的零和假设——现实中市场容量不是硬性的零和，而是弹性的', bold_prefix='')

add_h2('7.4 评价保真度：约 50%')
add_para('评价保真度是三者中最低的，因为"什么是好的经纪人表现"本身就是一个高度主观的判断。')
add_para('当前的三轴评分有一个深层假设：')
add_highlight('卖出价格是最重要的（ability 40 分），保住房源第二重要（defense 35 分），业主满意度第三重要（satisfaction 25 分）。')
add_para('但现实中，很多优秀经纪人的核心竞争力恰恰是满意度优先——"业主满意了自然会介绍新客户"。当前的评分体系没有考虑口碑的乘数效应（wordOfMouth 已计算但未接入评分）。')
add_para('此外，goalTier 权重的设计也反映了"哪些房源更重要"的主观判断。在现实中，经纪人可能因为一套 normal 房的漂亮成交而获得行业口碑——这种"无差别成交的价值"在当前评分中被 goalTier 权重压低了。')

# ============================================================
# 第八章
# ============================================================
add_h1('第八章 设计启示：世界模型理论给我们的三个洞见')

add_h2('8.1 洞见一：表征对齐（Representation Alignment）')
add_para('JEPA 的核心教义是：世界模型的输出表征应该是 cost module 的输入表征。如果两者不匹配，系统就会出现"信息漏斗"——精心维护的状态维度没有被评价函数使用，而评价函数需要的信息没有被状态维护。')
add_para('在王牌顾问中，存在以下表征不对齐：')
add_table(
    ['维护的状态', '评价函数使用', '对齐状态'],
    [
        ['wordOfMouth（口碑）', '终局评分未使用', '未对齐（管道已建，阀门未开）'],
        ['D1 中的 poolSize', 'ability 只看成交结果', '部分对齐（过程 vs 结果）'],
        ['promotionBudget 使用量', '评分刻意不使用', '有意不对齐（设计决策）'],
        ['customerRuntimeState.churnRisk', 'D1 未使用', '未对齐（D1 只看数量不看质量）'],
        ['advisorTrust（客户对经纪人信任）', 'D1 未使用', '未对齐（同上）'],
    ]
)
add_para('JEPA 的建议是：让表征空间对齐——如果 churnRisk 是重要的，就应该反映在竞争力评估中；如果 wordOfMouth 是有价值的，就应该接入评分体系。')

add_h2('8.2 洞见二：反事实模拟（Counterfactual Simulation）')
add_para('Dreamer 的核心优势是能在世界模型中"想象"不同行动的后果。王牌顾问的世界模型虽然支持前向推演，但缺乏：')
add_bullet('"如果我昨天没有维护那套房，今天会怎样"的反事实推理')
add_bullet('"如果我先推 A 房再推 B 房 vs 先推 B 房再推 A 房"的并行路径比较')
add_bullet('"如果竞品在 Day 5 发动攻势，我的防线在哪"的压力测试')
add_para('这些功能不是当前游戏的核心需求（玩家自己在脑中做反事实推理），但如果有 AI 辅助决策的需求（如教练模式），反事实模拟将变得至关重要。')
add_para('技术上，可以通过"分支世界"（branching world）的方式实现：在关键决策点 fork GameState，分别推演不同选择的后果，最后比较终局评分。')

add_h2('8.3 洞见三：非线性涌现（Nonlinear Emergence）')
add_para('当前的世界模型以线性公式为主（加权求和、乘法调制）。但真实世界的很多现象是非线性的：')
add_bullet('"信任崩溃"：信任从 50 到 40 是渐进的，但从 40 到 30 可能是突然的（"最后一根稻草"效应）', bold_prefix='')
add_bullet('"口碑爆发"：一套房的漂亮成交可能带来 3 套新房源，而非线性的 1 套', bold_prefix='')
add_bullet('"竞品雪崩"：当你的信任低于某个临界点，竞品可能同时发起攻势', bold_prefix='')
add_para('学习型世界模型（如 Dreamer 的 RSSM）天然能捕捉这些非线性效应，因为神经网络的激活函数可以表达任意非线性映射。手工世界模型则需要显式编码这些非线性关系——这在技术上可行（如使用 sigmoid 替代线性衰减），但会增加系统的复杂度和维护成本。')
add_highlight('一个折中方案是：在手工世界模型的框架内引入"临界点"机制——当某个状态变量越过阈值时，触发非线性效应。当前系统中已有类似设计（如 windowDays ≤ 0 时的续约判断），但可以进一步扩展。')

# ============================================================
# 第九章
# ============================================================
add_h1('第九章 总结：手工世界模型的定位与未来')

add_h2('9.1 我们在世界模型理论光谱上的位置')
add_para('将各种世界模型按"学习程度"排列在一个光谱上：')
add_code('纯手工 ◄──────────────────────────────────────────► 纯学习')
add_code('  │                                                    │')
add_code('  │  王牌顾问    │    混合系统    │    Dreamer    │    GPT')
add_code('  │  (全部手工)  │  (部分手工    │  (全部学习    │  (全部学习')
add_code('  │              │   部分学习)    │   + 环境交互)  │   + 海量数据)')
add_para('王牌顾问处于光谱的最左端——一个完全手工构建的世界模型。这不是一个弱点，而是一个设计选择，它带来了：')
add_bullet('极高的可解释性：每个公式都有业务含义，每个常量都可以被质疑和调整')
add_bullet('极快的迭代速度：改一行代码就能看到效果，不需要重新训练')
add_bullet('极低的计算成本：每天 24 步流水线 < 5ms，不需要 GPU')
add_bullet('极强的设计控制力：每个状态变量的含义都是明确的')

add_h2('9.2 手工世界模型的独特价值')
add_para('在 AI 研究的语境中，手工世界模型常被视为"落后的"或"需要被学习型替代的"。但王牌顾问的实践证明了另一种可能：')
add_highlight('手工世界模型是领域知识的精确编码——它不是"没有学到"，而是"已经知道"。')
add_para('120+ 个平衡常量中的每一个都是一个业务洞察的数值化表达。这些洞察不可能从有限的游戏数据中自动学习到——它们来自对上海二手房市场的深入理解、对经纪人工作模式的长期观察、以及对玩家心理的精心设计。')

add_h2('9.3 未来可能的演进方向')
add_para('如果要将王牌顾问的世界模型向学习型世界模型演进，以下是最自然的路径：')

add_h3('方向一：数据驱动的常量校准')
add_para('保留手工设计的公式结构，但用真实游戏数据（或大量模拟数据）来校准 120+ 个常量。这相当于用学习来补充手工设计的不足，而非替代手工设计的结构。')

add_h3('方向二：混合架构')
add_para('在手工世界模型的基础上，增加一个学习型的"修正模块"——当手工公式预测的 next state 与实际 next state 有偏差时，修正模块用一个小网络来补偿。这类似于"残差学习"（residual learning）。')

add_h3('方向三：Coach AI 的反事实引擎')
add_para('利用现有的 resolveOneDay 流水线作为"世界模型模拟器"，在多个分支世界上并行推演，为玩家的每个决策提供"如果你这样做了，结果会怎样"的反事实分析。这不需要改变世界模型本身，只需要增加一个"分支管理器"。')

add_h2('9.4 最终定位')
add_para('王牌顾问的世界模型在理论光谱上的定位是：')
add_highlight('一个高度结构化、完全可解释、领域知识密集的手工世界模型。它的精度受限于设计者的认知边界，但它的精确性和可维护性是学习型世界模型无法比拟的。')
add_para('它不是一个"简陋版 Dreamer"，而是一个"精确版领域模拟器"。它不追求"从数据中发现规律"，而是追求"用代码精确表达已知规律"。')
add_para('在 AI 研究中，手工世界模型正在被学习型世界模型超越。但在游戏设计中，手工世界模型仍然是最好的选择——因为游戏需要的不是"最真实的模拟"，而是"最有趣的模拟"。')
add_para('而"有趣"这个目标，到目前为止，仍然是人类设计者的直觉最可靠。')

doc.add_paragraph()
add_highlight('—— 全文完 ——')

# ============================================================
# 保存
# ============================================================
output_path = '/Users/jiaqi/Documents/开放日测算/王牌顾问-世界模型理论分析.docx'
doc.save(output_path)
print(f'Word document saved to: {output_path}')
