from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'PingFang SC'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'PingFang SC')

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hf = hs.font
    hf.name = 'PingFang SC'
    hf.bold = True
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'PingFang SC')

sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

def add_title(text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(26)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True

def add_h1(text):
    doc.add_heading(text, level=1)

def add_h2(text):
    doc.add_heading(text, level=2)

def add_h3(text):
    doc.add_heading(text, level=3)

def add_para(text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_highlight(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0x0d, 0x47, 0xa1)
    run.font.size = Pt(11)
    return p

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Menlo'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pf = p.paragraph_format
    pf.left_indent = Cm(1)
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    return table

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

# ============================================================
# 封面
# ============================================================
add_title('王牌顾问 · 系统设计鉴赏')
add_subtitle('—— 一场用代码还原的上海二手房经营世界')
doc.add_paragraph()
doc.add_paragraph()

# ============================================================
# 第一章
# ============================================================
add_h1('第一章 概览：这个世界在做什么')

add_h2('1.1 一句话定义')
add_para('王牌顾问模拟的是一名上海二手房经纪人在 21 天内经营 3~6 套房源的完整周期。玩家的核心命题只有一个：', bold_prefix='核心命题：')
add_para('把房子卖掉，卖得漂亮。', bold_prefix='')
add_para('但"卖掉"这件事在一个真实市场里从来不是经纪人一个人说了算的。世界围绕三方博弈构建——')
add_code('        经纪人（你）')
add_code('       ╱          ╲')
add_code('  业主（信任你的房东）  客户（买房的人）')
add_code('       ╲          ╱')
add_code('     竞品门店 & 竞品房源')
add_code('          │')
add_code('     公司内部压力')
add_code('          │')
add_code('     市场宏观周期')
add_para('你需要同时应对业主信任衰减、客户犹豫流失、竞品暗中撬房、公司持续施压、市场情绪起伏——每一天都在做取舍，每一个动作都在改变多个指标，每一局都因为随机种子而不同。')

add_h2('1.2 运行时三层架构')
add_para('系统严格遵循 DDD 四层架构，但在运行时可以更直观地理解为三层：')
add_table(
    ['层', '回答什么', '关键概念'],
    [
        ['World', '这个游戏世界里可能发生什么', '商圈、房源原型、客户库、业主原型、事件模板'],
        ['Scenario', '这一局如何开始、教什么、怎么结束', '难度配置、角色分配、剧本事件、竞争拓扑、目标分'],
        ['Run', '玩家这一把当前打到了哪里', '每日状态、引擎计算、动作执行、事件流'],
    ]
)
add_para('World 是只读的版本化内容；Scenario 是一局的蓝图，在开局时 structuredClone 固化到 RunContext 中，之后不可变；Run 才是那个每天都在变的可变状态。')
add_code('World (只读) ──初始化──▶ Scenario (只读快照) ──createInitialState──▶ Run (可变)')
add_code('                                                                    │')
add_code('                                              resolveOneDay(state) ◀─┘')
add_code('                                                    ↓')
add_code('                                              24 步引擎流水线')
add_code('                                                    ↓')
add_code('                                              新一天的 GameState')
add_highlight('亮点：World/Scenario/Run 三层分离让"同一个世界生成不同局"成为可能。不同的 seed + 不同的 difficultyId 会产出完全不同的房源组合、竞争组拓扑和剧本事件弧，但底层共享同一套市场单元格和客户原型——每一局都是"同一个上海"的不同侧面。')

add_h2('1.3 难度体系：6 档梯度的精细调控')
add_para('难度不只改一个数字，而是一套 17 维旋钮矩阵。每一档难度都在 OutcomeControlRules 中精细定义了玩家和竞品的相对能力：')
add_table(
    ['维度', 'warmup', 'easy', 'standard', 'advanced', 'hard', 'extreme'],
    [
        ['房源数', '3', '4', '5', '5', '6', '6'],
        ['天数', '21', '21', '21', '21', '21', '21'],
        ['市场成交容量(21d)', '5', '5', '4', '4', '3', '3'],
        ['玩家基础成交预期', '2', '2', '1', '0.68', '0.7', '0.66'],
        ['玩家成交签约倍率', '1.0', '1.0', '0.84', '0.72', '0.78', '0.76'],
        ['客户漏斗推进倍率', '1.08', '1.08', '1.0', '0.86', '0.85', '0.78'],
        ['竞品门店能力倍率', '0.75', '0.75', '1.08', '1.25', '1.25', '1.45'],
        ['竞品抢房倍率', '0.65', '0.65', '1.08', '1.28', '1.18', '1.32'],
        ['目标分', '58', '64', '72', '78', '84', '88'],
    ]
)
add_para('这组数字讲了一个清晰的梯度故事：')
add_bullet('warmup/easy：市场给 5 个成交名额，你至少能拿 2 个，竞品被压到 0.75 倍——你是在学机制，不是在拼命。', bold_prefix='入门：')
add_bullet('standard：名额缩到 4，你只能稳拿 1 个，竞品升到 1.08 倍——开始真正竞争了。', bold_prefix='标准：')
add_bullet('advanced：你的签约倍率跌到 0.72，竞品门店升到 1.25——"推到桌前"和"签下字来"都变难了。', bold_prefix='进阶：')
add_bullet('extreme：3 个名额、0.76 签约倍率、1.45 倍竞品——每个成交都是抢回来的。', bold_prefix='极难：')
add_para('难度还影响初始条件——extreme 的业主初始 trust 只有 46~58（对比 warmup 的 62~74），窗口天数只有 4~7 天（对比 warmup 的 10~16 天）。还没开始你就已经在悬崖边上了。')
add_highlight('亮点：playerBaseDealExpectation21d 可以是小数（如 0.68），系统用确定性的 FNV 哈希来决定这一局是 0 还是 1 个保底名额——同一个种子永远产生相同的结果，不同种子则可能差 1 个名额。这让 replay 有意义，也让"差一局"的感觉成为常态。')

add_h2('1.4 引擎全景与依赖关系')
add_para('系统包含 20+ 个引擎/子系统，在 resolveOneDay 中被编排为 24 步顺序执行。它们之间存在严格的读写依赖：')
add_code('                        ┌─────────────────────┐')
add_code('                        │  1 释放成交名额       │')
add_code('                        │  2 市场单元格漂移     │──写 marketPrice')
add_code('                        │  3 季节性修正        │     │')
add_code('                        └─────────┬───────────┘     │')
add_code('                                  ↓                  ↓')
add_code('                        ┌─────────────────────────────┐')
add_code('                        │  4-5 每日事件               │──写 demandHeat/sentiment')
add_code('                        │  6-7 竞品门店/房源波动      │──写 rivalListings[]')
add_code('                        │  8 竞品压力传导            │──写 case.heat/trust')
add_code('                        └─────────┬─────────────────┘')
add_code('                                  ↓')
add_code('                   ┌──────────────────────────────┐')
add_code('                   │  9-10 公司压力               │──写 customerStates')
add_code('                   │  11 客户属性波动             │')
add_code('                   │  12 客户运行时推进           │──写 interest/confidence/stage')
add_code('                   │  13 竞品拉力影响客户         │──写 churnRisk')
add_code('                   └──────────┬───────────────────┘')
add_code('                              ↓')
add_code('              ┌─────────────────────────────────┐')
add_code('              │  14 机会漏斗推进                 │──写 intent/confidence/stageIndex')
add_code('              │  15 客户反馈→房源               │──写 case.heat/trust/viewings')
add_code('              │  16 竞争组+抢房判定             │──写 case.status')
add_code('              └──────────┬──────────────────────┘')
add_code('                         ↓')
add_code('          ┌──────────────────────────────────┐')
add_code('          │  17 剧本事件触发                  │──写 trust/urgency/heat/windowDays')
add_code('          │  18 成交结算                      │──写 case.status=sold + 全链路副作用')
add_code('          │  19 竞品争夺剩余名额              │──写 rivalClaimedDeals')
add_code('          │  20 房源自然衰减                  │──读 ownerArchetype 写 trust/heat/urgency')
add_code('          └──────────┬───────────────────────┘')
add_code('                     ↓')
add_code('          ┌──────────────────────────────────┐')
add_code('          │  21 被动线索生成                  │──创建新 opportunity')
add_code('          │  22 随机事件                      │──全局冲击')
add_code('          │  23 信号衰减                      │')
add_code('          │  24 周度复盘+资金+聚焦会+新一天   │')
add_code('          └──────────────────────────────────┘')
add_para('核心因果链：环境先变 → 人再反应 → 环境对人做反馈 → 房源自身衰减 → 新线索作为明天的种子。')
add_highlight('亮点：updateDerivedState（调用 scoring.ts 重算 competitiveness）不在每步之间调用，而是在 resolveOneDay 末尾只调用两次。这是性能/一致性权衡——每天算 20 次评分既贵又不稳定，让引擎在"昨天的竞争力评估"上工作反而更可预测。')

# ============================================================
# 第二章
# ============================================================
add_h1('第二章 世界模型：上海二手房的数值还原')

add_h2('2.1 市场单元格——两座永不同步的城')
add_para('世界包含 2 个市场单元格，代表两个真实商圈：')
add_table(
    ['单元格', '商圈', 'demandHeat', 'supplyPressure', 'competitivePressure', 'sentiment'],
    [
        ['mc-qiantan', '浦东前滩 · 80-90㎡ 刚改', '76', '62', '58', '68'],
        ['mc-jingan', '静安寺北 · 60-75㎡ 资产型', '81', '54', '51', '71'],
    ]
)
add_para('4 个动态指标每日更新，核心驱动是一个正弦波函数：')
add_code('const wave = Math.sin((dayOfMonth + index * 4) / 8.5);')
add_code('demandHeat = clamp(demandHeat + wave * 3 + randomInt(-2, 2), 35, 92);')
add_para('index * 4 是关键——两个商圈的波峰错开约 4 天。当前滩在 Day 5 到达需求峰值时，静安还在低谷；等到 Day 9 静安升温时，前滩已经开始回落。这意味着玩家永远有事情要忙，但忙的方向在变。')
add_para('叠加在正弦波上的是每月的季节性因素（monthlyFactors[month] × seasonalityImpact）。3~5 月金三银四加 demandHeat，8~10 月秋季行情加 sentiment，12 月则回落——完全贴合上海二手房的实际季节规律。')
add_para('每套房源的 marketPrice 每日随所在单元格的供需差浮动：', bold_prefix='市场价传导：')
add_code('caseItem.marketPrice = max(')
add_code('  round(上日值 + (demandHeat - supplyPressure) / 18 + randomInt(-3, 3)),')
add_code('  round(askPrice * 0.84)    // 地板价 = 报价 × 84%')
add_code(')')
add_para('18 的除数意味着：demandHeat 比 supplyPressure 每高出 18 点，市场价每天涨 1 万。微弱但持续——如果你持续关注市场面板，会发现"今天市场价涨了 2 千"这种细节，而这 2 千可能正好让报价从"偏高"变成"严重偏高"，触发业主信任惩罚。')
add_highlight('亮点：两个商圈的异步脉动不是一个装饰，它是"资源分配"决策的信息基础。前滩需求高时应该把推广金投到前滩房源上，但如果此时静安业主的信任在崩溃，你又不能完全放弃——这种两难正是真实经营的常态。')

add_h2('2.2 业主原型——四种人格，四种衰减速度')
add_para('4 种业主原型是这个世界里最精妙的人格设计。它们不是标签，而是数值修正器：')
add_table(
    ['原型', 'trustDecayMultiplier', 'priceElasticity', 'urgencyGrowthBonus', 'patienceDelta', '推荐策略'],
    [
        ['焦虑型', '1.4', '1.25', '0.9', '-2', 'deep-cut（快刀斩乱麻）'],
        ['等价型', '0.9', '0.7', '0.6', '+1', 'hold-story（讲好故事慢慢来）'],
        ['试水型', '1.15', '0.95', '1.1', '-3', 'small-cut（小步快跑）'],
        ['博弈型', '1.05', '0.85', '1.2', '0', 'hold-story（信息博弈）'],
    ]
)
add_para('trustDecayMultiplier 的 55% 跨度是整个系统的核心张力来源。同样是"今天没碰业主"（基础 trust 损失 2.8 点）：', bold_prefix='关键洞察：')
add_bullet('焦虑型业主损失 2.8 × 1.4 = 3.9 点')
add_bullet('等价型业主损失 2.8 × 0.9 = 2.5 点')
add_para('从 trust=76 跌到 trust=50（从"可能满意"到"开始后悔"）：')
add_bullet('焦虑型：26 ÷ 3.9 ≈ 7 天')
add_bullet('等价型：26 ÷ 2.5 ≈ 10 天')
add_para('3 天的差距，在 21 天的局里几乎是 15% 的经营时间。')
add_para('当报价超过市场价 5% 时，系统计算超价惩罚：', bold_prefix='priceElasticity 同样深藏杀机：')
add_code('caseItem.trust -= overpricedElasticityBasePenalty + max(0, ownerArchetype.priceElasticity - 1);')
add_para('焦虑型的 priceElasticity=1.25，额外惩罚 0.25 点/天；等价型的 priceElasticity=0.7，不触发额外惩罚。')
add_highlight('亮点：四种原型不是"性格标签"，而是一套持续消耗的乘数系统。你在 Day 1 看到焦虑型业主 trust=70 时，知道的是"只有 7 天安全期，之后每一天都是倒计时"。这种隐含的倒计时是所有紧张感的来源。')

add_h2('2.3 房源——一台精密的状态机')
add_para('每套房源是一个拥有 30+ 字段的复杂状态机。核心属性构成三轴信号：')
add_para('D1 漏斗健康度——"有没有人在看这套房"：', bold_prefix='')
add_bullet('poolSize：总线索数')
add_bullet('activeContacts：活跃客户数')
add_bullet('funnelWeight：晚期阶段权重')
add_bullet('advanceSpeed：阶段升格速度')
add_bullet('stagnationRisk：停滞风险')
add_para('D2 房源禀赋——"这套房本身好不好"：', bold_prefix='')
add_bullet('7 轴打分：layout(0.2) / neighborhood(0.2) / amenity(0.15) / decor(0.15) / light(0.1) / floor(0.1) / structure(0.1)')
add_bullet('从 housePrototype.axisScores 初始化，整个局不再变化——禀赋是天生的，你改不了。')
add_para('D3 业主配合度——"业主愿不愿意配合你"：', bold_prefix='')
add_bullet('priceFlex / patience / urgency / trust / consistency / baseline')
add_bullet('recentCooperation = 3 天内有动作 → +25')
add_bullet('consistency = 最近 7 天动作天数 / 7 × 40 + 30')
add_para('三轴合成竞争力：', bold_prefix='')
add_code('competitiveness = d1 × 0.5 + d2 × 0.25 + d3 × 0.25')
add_para('客户线权重 50%，禀赋 25%，意愿 25%——这个权重分配本身就是一个设计宣言：经纪人的核心工作是"把人带到桌前来"，不是改造房子。')
add_para('tickCases 中处理 6 条独立的衰减路径：', bold_prefix='每日自然衰减：')
add_bullet('不碰业主 → trust 衰减 × ownerArchetype.trustDecayMultiplier', bold_prefix='1. ')
add_bullet('不碰房源 → heat 衰减；emotional 业主额外 × heatSensitivity', bold_prefix='2. ')
add_bullet('pragmatic 业主 → 对价格偏差敏感：紧凑价 +2 trust，宽松价 -2 trust', bold_prefix='3. ')
add_bullet('报价过高 → 三重打击：trust / heat / patience 各减，乘以 priceElasticity', bold_prefix='4. ')
add_bullet('emotional 业主 + 低热度 → 额外 trust 损失 × heatSensitivity', bold_prefix='5. ')
add_bullet('urgency 自然增长 → urgent 固定 +5，其他 +2~3，短窗口（≤6 天）额外 +2', bold_prefix='6. ')
add_para('当 windowDays 降到 0 时，检查 trust ≥ 76 && satisfaction ≠ unhappy && d3 ≥ 62。满足则续 4 天窗口但 trust -6。这是一个"抢救"机制——业主还在给你机会，但信任在消耗，而且续约的 trust 扣除会让下一次续约更难。', bold_prefix='续约窗口：')
add_highlight('亮点：D1 的 log2 边际递减让"从 0 到 1"的客户获取比"从 9 到 10"更有价值——鼓励玩家把资源分散到多套房源上，而不是只堆一套。')

add_h2('2.4 客户——一个双层数据模型')
add_para('客户是系统中最复杂的实体，因为它维护了两个独立的视角：')
add_bullet('Opportunity（经纪人视角）：我有一条来自"小红书"的线索，客户"张磊"，对"前滩华庭"感兴趣，intent=72，confidence=58，stageIndex=3。')
add_bullet('CustomerRuntimeState（客户视角）：张磊在比较"前滩华庭"和"静安雅苑"两套房，对前滩 interest=75、confidence=60，对静安 interest=68、confidence=55，churnRisk=32，fatigue=18。')
add_para('为什么需要双层？因为一个客户可以同时看多套房源。Opportunity 只记录"这条线索进展如何"，CustomerRuntimeState 才记录"这个客户的内心世界"。')
add_para('从静态属性到行为模式：', bold_prefix='决策风格推导：')
add_code('urgency ≥ 76 && activity ≥ 72 → decisive    // 阶段推进阈值 64')
add_code('priceSensitivity ≥ 72 || activity ≤ 54 → hesitant  // 阈值 78')
add_code('else → balanced                              // 阈值 70')
add_para('"果断型"客户在 interest=64 时就愿意推进阶段，"犹豫型"要到 78——同样的漏斗推进需要 14 点更多 intent。14 点可能就是两到三天不带看的差距。')
add_para('每个客户 × 每套关注的房源，每日演化公式：', bold_prefix='兴趣/信心每日演化：')
add_code('interest += (heat-55)/10 + (advisorTrust-50)/12 + interactionBoost')
add_code('         - fatiguePenalty - comparePenalty - rivalryPenalty + priceAdvantage + randomInt(-4,4)')
add_code('confidence += (trust-55)/14 + (d3-50)/16 - priceSensitivity/80×stagnation - rivalryPenalty×0.7 + randomInt(-3,3)')
add_para('当客户对两套房的 interest+confidence 差距 ≤ 10 时触发 comparing 状态。进入 comparing 的客户每日 +3 churnRisk，且未选中的那套房 interest -4 / confidence -3。这是竞品效应在客户侧的直接体现。', bold_prefix='comparing 状态：')
add_para('客户引擎算完后，反向同步到 opportunity 对象。这是整个系统里最关键的一致性保证。', bold_prefix='双向同步：')
add_highlight('亮点：d1 影响 intent，d3 影响 confidence——这个分配不是随意的。D1（漏斗健康度）衡量"有多少人在看"，直接影响客户的购买意愿；D3（业主配合度）衡量"业主好不好打交道"，直接影响客户的购买信心。一个客户可能"很想买但不敢买"——这对应着"D1 好 D3 差"。这是整个系统里最深层的业务映射。')

add_h2('2.5 竞品体系——双层施压架构')
add_para('竞品分两层，每层有不同的施压方式：')
add_para('3 种风格门店的 activityHeat 增量倍率不同：aggressive ×2, steady/traffic ×1, relationship ×0.5。门店不直接跟玩家竞争，而是为竞品房源提供能力乘数。', bold_prefix='竞品门店——持续性的环境压力：')
add_para('3 种原型：', bold_prefix='竞品房源——点对点的直接竞争：')
add_table(
    ['原型', 'baseHeat', 'leadSiphonPower', 'ownerAnchorPower', '含义'],
    [
        ['急售平替盘', '68', '58', '62', '低价走量，撬客户'],
        ['同公司重点盘', '60', '50', '42', '内部竞争，分资源'],
        ['强卖点竞品', '64', '54', '50', '故事好，两头撬'],
    ]
)
add_para('这是竞品影响房源的核心机制：', bold_prefix='5 级压力传导模型：')
add_code('priceOverlap = clamp(1 - priceGap × 6, 0, 1);  // 价格差距 10% 以内才产生压力')
add_code('nearestPressure = Σ(priceOverlap × (leadSiphonPower + ownerAnchorPower + heat) / 3) / rivals.length;')
add_code('if (adjustedPressure >= 34) {  // 压力阈值 34')
add_code('  case.heat -= pressure / 100 × rivalPressureHeatImpact;')
add_code('  case.trust -= pressure / 100 × rivalPressureTrustImpact;')
add_code('  opportunity.intent -= pressure / 85;')
add_code('  opportunity.confidence -= pressure / 110;')
add_code('}')
add_para('34 的阈值制造了一个"温水煮青蛙"效应——一两套竞品出现时你感觉不到什么，但当竞品房源变多、价格重合度提高，压力越过 34 时，你的 trust/heat/intent 同时开始被侵蚀。')
add_highlight('亮点：极端难度的 rivalPressureHeatImpact=2.6、rivalPressureTrustImpact=1.3——是 standard 的 2.4 倍和 3.1 倍。同一套竞品房源在 extreme 难度下的侵蚀力度是 standard 的两三倍。')

add_h2('2.6 信号与事件——三层信息架构')
add_para('世界提供三层信息，从弱到强：')
add_para('3 种模板（buyer_demand / seller_intent / rival_activity），每日 18% 概率生成，有 expiresInDays 衰减。信号不直接影响任何数值，只给玩家提供决策情报。', bold_prefix='第一层：市场信号——')
add_para('6 种类型，加权随机。signal_only 权重 4（纯信息），其余权重 2（有具体后果）。每日事件的概率由难度决定：warmup 10%，extreme 38%。', bold_prefix='第二层：每日事件——')
add_para('3 种宏观冲击：', bold_prefix='第三层：随机事件——')
add_bullet('policy-shift：所有活跃客户 confidence -10（"限购收紧了"）')
add_bullet('school-boom：随机商圈 heat +18, sentiment +12（"学区利好"）')
add_bullet('competitor-cut：所有商圈 competitivePressure +18，活跃房源 heat -4（"竞品集体降价"）')
add_para('随机事件的权重也随难度翻转：warmup 中 school-boom 权重 5、competitor-cut 权重 1；extreme 中 school-boom 权重 1、competitor-cut 权重 7。')
add_highlight('亮点：三层信息的层级设计不是功能分解，而是认知负荷管理。新手只看信号就能做决策；老手会结合每日事件预判趋势；高手需要在随机事件的冲击下做应急预案。')

add_h2('2.7 场景生成——从蓝图到实例')
add_para('每一局由 assembleGeneratedScenario 从 3 个维度组装：')
add_bullet('难度配置（DifficultyProfile）：6 档，每档 17+ 维旋钮', bold_prefix='1. ')
add_bullet('蓝图（ScenarioBlueprint）：8 个预置蓝图，决定角色槽位、竞争拓扑、事件弧', bold_prefix='2. ')
add_bullet('种子（seed）：确定性随机源，同一 seed 产出同一局', bold_prefix='3. ')
add_para('每套房源被分配一个角色：', bold_prefix='角色槽位：')
add_table(
    ['角色', '含义', 'goalTier'],
    [
        ['anchor', '本局核心盘，绝对不能丢', 'core'],
        ['fragile', '脆弱盘，容易崩', 'important'],
        ['traffic', '流量盘，客源发动机', 'important'],
        ['grind', '苦活盘，硬啃才有结果', 'normal'],
        ['spoiler', '干扰盘，让别人难做', 'normal'],
        ['sacrifice', '牺牲盘，放弃也不心疼', 'normal'],
    ]
)
add_para('3 种竞争组结构：', bold_prefix='竞争拓扑：')
add_bullet('district_clusters：同商圈所有房源一个组，价格弹性 0.86~1.08')
add_bullet('paired_pressure：每个商圈选 2~3 套房配对，价格弹性 0.92~1.16')
add_bullet('chain_clusters：同商圈全入组且高弹性，价格弹性 1.08~1.33')
add_para('6 种叙事线（从 relationship_recovery 到 competition_collapse），每种 2~3 个剧本事件，Day 偏移使用 ±1 天的随机抖动。', bold_prefix='事件弧：')
add_highlight('亮点：eventDay 偏移的 ±1 天随机抖动打破了确定性剧本的机械感，让同一蓝图在不同种子下产生微妙不同的节奏。')

# ============================================================
# 第三章
# ============================================================
add_h1('第三章 资源体系：时间、金钱和注意力的三重约束')

add_h2('3.1 精力——不可存储的时间压力')
add_para('精力每日重置，由 WEEKLY_ROUTINE 决定上限：')
add_table(
    ['星期', '精力', '主题', '设计意图'],
    [
        ['周一', '6', '客户拓展', '一周之始，适合拉新'],
        ['周二', '5', '房源维护', '趁热打铁稳业主'],
        ['周三', '5', '带看推进', '推漏斗的关键日'],
        ['周四', '5', '聚焦会', '提报+评选，资源杠杆'],
        ['周五', '7', '周度冲量', '能量最多，多线操作'],
        ['周六', '3', '客户沉淀', '能量最少，只能选重点'],
        ['周日', '3', '复盘调整', '选策略，调节奏'],
    ]
)
add_para('每周总能量 = 34 点。6 套房源各需 1 点碰触 + 1 点业主沟通 = 每天至少 12 点。但每天最多 7 点——你不可能同时照顾所有房源。')
add_para('精力不可存储——今天不用就浪费了。这制造了"每天都有必须做的事"的时间压力。周末 3 点精力时，玩家被迫做取舍。')
add_highlight('亮点：周末 3 点精力的设计不是"让玩家休息"，而是"强迫玩家暴露优先级"。当你只能在 6 套房中选 3 套维护时，你的选择就暴露了你的判断——判断对了加分，判断错了可能丢房。')

add_h2('3.2 推广金——可累积的运营资本')
add_para('推广金是累计制资源，通过 budget.ts 的全量流水记账。')
add_para('周度拨付（weekly-allocation）和成交返还（sale-rebate：佣金 × promotionRebateRatio，最低 promotionRebateFloor）。', bold_prefix='来源：')
add_para('部分动作需要推广金（小红书推广、经纪人投放、私域转介绍）。', bold_prefix='消耗：')
add_para('成交 → 佣金 → 推广金返还 → 更多推广 → 更多客户 → 更多成交。但这个正反馈被 marketDealSlots 机制截断——市场每天的名额有限。', bold_prefix='正反馈闭环：')
add_highlight('亮点：推广金"只解释资源打法，不直接代表这局打得好"——这句话出现在 resultEvaluation.ts 的 coachNotes 中。系统刻意不把推广金消耗量纳入评分，防止"多花钱=高分"的误导。')

add_h2('3.3 周四聚焦会——每周一次的杠杆')
add_para('聚焦会是整个资源体系里最精巧的子系统：')
add_para('玩家提报最多 3 套房（消耗 1 精力）→ 系统用 focusMeetingScore 排序选前 2 套 → 入选房源获得 heat+12, trust+4, patience+3, isFocused=true', bold_prefix='流程：')
add_code('focusMeetingScore = stageBonus          // lead.stageIndex ≥ 4 → 180, ≥ 3 → 120, ≥ 2 → 70, else 30')
add_code('                  + quoteBonus           // intent ≥ 70 → 36, ≥ 55 → 24, else 10')
add_code('                  + competitiveness × 1.3')
add_code('                  + heat × 0.9')
add_code('                  + pressureBonus        // max(0, 8 - windowDays) × 8')
add_code('                  + offers × 16')
add_para('被选中的房源在后续的 spawnPassiveLeads 中获得 passiveLeadFocusedMultiplier 加成。这形成了一个聚焦飞轮：提报 → 入选 → 流量加成 → 更多客户 → 更容易成交 → 下周更容易入选。')
add_highlight('亮点：pressureBonus = max(0, 8 - windowDays) × 8——窗口只剩 3 天时 bonus=40，只剩 1 天时 bonus=56。这意味着危急房源更容易被聚焦会选中。系统在帮你——你最该关注的那套房，评分公式也在推它。')

add_h2('3.4 口碑——隐藏的长期资产')
add_para('口碑的累积规则：成交 +4（base）+ 策略加成（hold:0 / balanced:+1 / close:+2），核销 -3。口碑目前完全不影响任何引擎行为和评分计算，只在 auxiliaryStats 中累积。它是一个设计预留。')
add_highlight('亮点：close 策略的 wordOfMouthBonus=2 比 hold 的 0 高，这意味着"快速成交"在口碑维度上有额外收益。如果未来口碑接入评分，close 策略将获得一个平衡其价格劣势的新维度。')

# ============================================================
# 第四章
# ============================================================
add_h1('第四章 引擎与计算逻辑：24 步流水线的精密齿轮')

add_h2('4.1 每日结算流水线')
add_para('resolveOneDay 是整个系统的心脏，24 步严格顺序执行。每一步都是 (state) => void 的就地修改函数。')
add_table(
    ['步', '函数', '读', '写', '关键阈值'],
    [
        ['1', 'releaseMarketDealSlotsForDay', 'slotSchedule', 'releasedSlots', '—'],
        ['2', 'updateMarkets', 'day, rngState', 'markets[].4指标, marketPrice', 'demand∈[35,92]'],
        ['3', 'tickSeasonality', 'month', 'demandHeat/sentiment', '—'],
        ['4-5', 'roll/applyDailyMarketEvent', 'dailyEventPool', 'markets[]/rivalListings[]/customers[]', '—'],
        ['6', 'tickRivalStores', 'rivalStores[]', 'activityHeat', '—'],
        ['7', 'tickRivalListings', 'rivalListings[]', 'freshness/heat/status', '—'],
        ['8', 'applyRivalPressure', 'rivalListings[], cases[]', 'heat/trust, intent/conf', 'threshold=34'],
        ['9', 'tickCompanyPressure', 'companyPressure', 'companyPressure.*', '—'],
        ['10', 'applyCompanyPressure', 'companyPressure, customers[]', 'customerStates[]', 'sharedLead≥58'],
        ['11', 'updateCustomers', 'customers[]', 'activity/urgency', '—'],
        ['12', 'progressCustomerDemand', 'customers[], cases[]', 'interest/conf/stage', '—'],
        ['13', 'applyRivalPullOnCustomers', 'rivalListings[], customers[]', 'interest/churnRisk', '—'],
        ['14', 'tickOpportunities', 'cases[], opportunities[]', 'intent/conf/stageIndex', 'intent≥82升格'],
        ['15', 'applyCustomerFeedbackToCases', 'customerStates[], cases[]', 'heat/trust/viewings', '—'],
        ['16', 'tickCompetition', 'competitionGroups[], cases[]', 'heat/trust/status', 'prob∈[0.5%,18%]'],
        ['17', 'fireScheduledEvents', 'scriptedEvents[]', 'trust/urgency/heat/window', '—'],
        ['18', 'settlePendingDealClosings', 'opportunities[], cases[]', 'status=sold+全链路', 'prob∈[0,95%]'],
        ['19', 'tryClaimOpenMarketDealForRivals', 'marketOutcome, day', 'rivalClaimedDeals', 'day≥maxDay-7'],
        ['20', 'tickCases', 'cases[], ownerArchetype', 'trust/heat/urgency/patience', 'trust≥76续约'],
        ['21', 'spawnPassiveLeads', 'cases[], markets[]', '新建 opportunities[]', '—'],
        ['22', 'triggerRandomEvent', 'randomEventPool', 'markets[]/customers[]/cases[]', '—'],
        ['23', 'settleMarketSignals', 'marketSignals[]', 'expiresInDays', '—'],
        ['24', '周度复盘+资金+聚焦会+新一天', '—', 'budget, focusMeeting, day++', '—'],
    ]
)

add_h2('4.2 机会引擎——从线索到成交的漏斗')
add_para('每条新线索出生时的质量公式：', bold_prefix='机会创建：')
add_code('intent = clamp(46 + bonus + fit×0.24 + heat×0.14 + activity×0.12 + channelQuality×10 - pricePenalty, 35, 89)')
add_code('confidence = clamp(48 + fit×0.25 + trust×0.16, 30, 92)')
add_para('权重排序：fit(0.24) > trust(0.16) > heat(0.14) > activity(0.12)。匹配度是线索质量的第一决定因素。')
add_para('', bold_prefix='机会每日推进：')
add_code('intent += (heat-55)/10 + (d1-50)/16 + randomInt(-4,4) - pricePenalty')
add_code('confidence += (d3-50)/14 + randomInt(-3,3)')
add_code('untouched → intent -= 4;  stagnationTicks++')
add_para('阶段升格门槛：intent >= 82 && chance(0.35 × playerFunnelProgressionScale)。82 的阈值 + 0.35 的概率 = 平均需要 intent 在 82 以上持续约 3 天才能升格。')
add_highlight('亮点：untouched → intent -= 4——4 天不碰一条线索，intent 掉 16 点，从 82 直接跌到 66，远低于升格阈值。这是"你不去维护，线索就会死"的直接数值表达。')

add_h2('4.3 竞品抢房引擎——6 种漏洞的精密判定')
add_para('shouldLoseToRival 是全系统最精细的风险模型。"业主不满 + 竞品有力 + 你没防住"三条件同时成立才触发。')
add_table(
    ['#', '漏洞类型', '触发条件（visibleSlip）', '竞品利用条件（rivalHasOpening）'],
    [
        ['1', 'urgentOpening', 'window ≤ 1 天 OR shadow 线索 ≥ 2', '—'],
        ['2', 'relationshipOpening', '≥ 4 天没碰业主 AND trust ≤ 58', '—'],
        ['3', 'trustCollapse', 'trust ≤ 36', '—'],
        ['4', 'coldAndNeglected', 'heat ≤ 24 AND ≥ 3 天没碰业主', '—'],
        ['5', 'pipelineOpening', '无活跃线索 OR 无合格线索+竞争压力高', '—'],
        ['6', 'priceAndPressureTrap', '价格/竞争双高', 'trust≤48 OR ≥3天没碰 OR window≤2'],
    ]
)
add_code('baseProb = Σ(visibleSlip[i] AND rivalHasOpening[i]) / validCount')
add_code('adjustedProb = baseProb × (1 + group.customerSpillover × 0.3) × rivalLossProbabilityScale')
add_code('finalProb = clamp(adjustedProb, 0.005, 0.18)  // 0.5% ~ 18%')
add_para('recentlyMaintained（最近 2 天内有维护动作）可以 guardAgainst 大部分漏洞，但当 pipelineOpening 存在时防护效果只有 60%。')
add_para('rivalLossProbabilityScale 从 warmup 的 0.08 到 extreme 的 0.9——11 倍差距。')
add_highlight('亮点：threatCooldownDays = 3——被撬一次后有 3 天冷却期，防止连续被撬到心态崩。但 3 天也意味着如果你不立刻补维护，第 4 天危机就回来了——"给你窗口但不给你安全感"。')

add_h2('4.4 成交引擎——两日结算的精密天平')
add_para('成交不是当天完成的，而是跨越两天的异步结算：今天报价 → 明天结算。')
add_table(
    ['策略', 'priceFactor', 'shift', 'loss', 'wordOfMouthBonus', '含义'],
    [
        ['hold', '1.0', '-6', '14', '0', '坚守报价，概率低但价格好'],
        ['balanced', '0.99', '+4', '8', '+1', '适度让利，性价比最高'],
        ['close', '0.985', '+9', '5', '+2', '快速成交，价格最低但概率最高'],
    ]
)
add_para('', bold_prefix='成交概率计算：')
add_code('successScore = intent×0.34 + confidence×0.26 + trust×(urgent?0.25:0.18) + competitiveness×0.16 - askPricePenalty×0.6 + strategy.shift')
add_code('probability = clamp(successScore × playerDealClosingScale, 0, 95)')
add_para('成交概率被 clamp 在 0~95%，永远不会是 100%。')
add_para('成交后的全链路副作用（finalizeClosedDeal）：', bold_prefix='成交后的连锁反应：')
add_bullet('房源状态 → sold，trust +8, heat +6')
add_bullet('调用 markCaseSold 写入所有 ending 字段')
add_bullet('佣金计算 + 推广金返还')
add_bullet('wordOfMouth +4 + 策略加成')
add_bullet('关闭该房源所有活跃机会')
add_bullet('客户侧：成交客户 → converted，其余 → idle/lost')
add_bullet('生成 closedDealRecord（含市场快照和价格快照）')
add_bullet('检查 marketDealSlots 配额')
add_bullet('触发 case_sold 事件')
add_highlight('亮点：成交概率永远不达 100%。即使 intent=100、confidence=100、trust=100，最终概率也只有约 85%。这 5% 的失败可能——"什么都对了但就是没签下来"——正是真实经营的常态。')

add_h2('4.5 结果判定引擎——8 种结局的中心化策略')
add_para('caseOutcome.ts 集中管理所有房源的结局判定，3 个入口函数覆盖 3 种终态。')
add_para('', bold_prefix='markCaseSold（成交时）：')
add_code('ownerSatisfaction: trust≥76 && soldPrice≥marketPrice×0.97 → happy')
add_code('                  trust≥62 → neutral, else → regret')
add_code('relativeOutcome: soldPrice/marketPrice ≥ 0.985 → outrun')
add_code('                soldPrice/marketPrice ≥ 0.95 → flat, else → lose')
add_para('', bold_prefix='markCaseWithdrawn（撤盘时）：')
add_code('ownerSatisfaction: trust≤50 → unhappy, else → regret')
add_para('', bold_prefix='markCaseLostToRival（被竞品抢走时）：')
add_code('ownerSatisfaction: trust≤52 → unhappy, else → regret  // 52 vs 50: 被抢走更容易不满')
add_table(
    ['endingType', 'bucket', '含义'],
    [
        ['sold_by_you_happy', 'good', '卖掉了，很满意'],
        ['sold_by_you_neutral', 'good', '卖掉了，无感'],
        ['not_sold_no_regret', 'good', '没卖掉，但不后悔'],
        ['switch_to_rent_no_regret', 'good', '转租也能接受'],
        ['sold_by_you_regret', 'neutral', '卖掉了，但体验不好'],
        ['not_sold_regret', 'neutral', '没卖掉，开始后悔'],
        ['sold_by_other', 'bad', '被别家卖掉了'],
        ['withdrawn_unhappy', 'bad', '彻底做崩了'],
    ]
)
add_highlight('亮点：lost_to_rival 的 unhappy 阈值是 52，withdrawn 的 unhappy 阈值是 50——差 2 点。被竞品抢走时业主的不满阈值更低，因为"被别人卖掉了"比"我自己撤盘"更让人恼火。')

add_h2('4.6 评分引擎——三轴加权与 goalTier 优先级')
add_para('终局评分由 evaluateFinalResult 在局末执行，三轴满分固定 40/35/25 = 总分 100。')
add_para('', bold_prefix='Ability 维度（满分 40）：')
add_code('base = outrun ? 1 : flat ? 0.65 : 0.2')
add_code('finishModifier = good ? 1.05 : bad ? 0.75 : 1')
add_code('score = clamp(Σ(normalizeGoalTier × base × finishModifier) / totalWeight × 40, 0, 40)')
add_para('', bold_prefix='Defense 维度（满分 35）：')
add_code('value = held ? 1 : at_risk ? 0.45 : 0')
add_code('score = clamp(Σ(normalizeGoalTier × value) / totalWeight × 35, 0, 35)')
add_para('', bold_prefix='Satisfaction 维度（满分 25）：')
add_code('value = happy:1 / neutral:0.75 / no_regret:0.65 / regret:0.3 / unhappy:0')
add_code('score = clamp(Σ(normalizeGoalTier × value) / totalWeight × 25, 0, 25)')
add_para('goalTier 权重：core=1.0 / important=0.7 / normal=0.4。核心房源在评分中的权重是 normal 房的 2.5 倍。')
add_table(
    ['等级', '条件', '标题'],
    [
        ['王牌', 'score ≥ ace', '这局你真的把房子卖顺了'],
        ['漂亮', 'score ≥ strong', '这局基本是你在带着节奏走'],
        ['过线', 'score ≥ pass', '至少把最关键的部分撑住了'],
        ['没保住', 'score < pass', '这局先交了一笔学费'],
    ]
)
add_highlight('亮点：scoreThresholds 推导中 ace = min(98, targetScore + 20)。extreme 难度 targetScore=88，ace=98；hard targetScore=84，ace=98——两者的 ace 线都是 98，区分度消失。这是一个已知的校准问题。')

# ============================================================
# 第五章
# ============================================================
add_h1('第五章 引擎与模型的关系：数据流、不变量与设计哲学')

add_h2('5.1 数据流全景')
add_para('系统的核心数据流可以从 World Spec → GameState → 引擎计算 → 评分 的链路来理解。GameState 是唯一的可变状态中心，所有引擎都是对它的就地修改。')

add_h2('5.2 引擎间的读写矩阵')
add_table(
    ['引擎', '读', '写'],
    [
        ['updateMarkets', 'day, rngState', 'markets[].4指标, cases[].marketPrice'],
        ['tickCases', 'day, ownerArchetype, cases[]', 'cases[].trust/heat/urgency/patience/windowDays'],
        ['tickOpportunities', 'day, cases[].heat/d1/d3/trust', 'opportunities[].intent/confidence/stageIndex'],
        ['progressCustomerDemand', 'day, customers[], cases[], opportunities[]', 'customerStates[].interest/confidence/stage, opportunities[](sync)'],
        ['applyRivalPressure', 'rivalListings[], cases[]', 'cases[].heat/trust, opportunities[].intent/confidence'],
        ['tickCompetition', 'competitionGroups[], cases[], rivalListings[]', 'cases[].heat/trust/status, rivalListings[]'],
        ['settlePendingDealClosings', 'opportunities[], cases[], marketDealSlots', 'cases[].status/soldPrice/ending字段, closedDeals[], customerStates[]'],
        ['applyCompanyPressure', 'companyPressure, customerStates[]', 'customerStates[].intent/confidence（shadow 客户）'],
    ]
)

add_h2('5.3 五大不变量')
add_para('系统通过 5 个不变量保护状态一致性：')
add_bullet('一套房最多成交一次：closedDeals 中 caseId 唯一', bold_prefix='1. ')
add_bullet('活跃机会单房源不超过 4 个：MAX_ACTIVE_OPPORTUNITIES_PER_CASE = 4', bold_prefix='2. ')
add_bullet('市场成交名额每日有限：marketDealSlotsPerDay 硬上限', bold_prefix='3. ')
add_bullet('信任/热度下限 10：clamp(trust, 10, 100)，房源不会完全归零', bold_prefix='4. ')
add_bullet('竞品抢房冷却期：threatCooldownDays = 3', bold_prefix='5. ')
add_para('还有 collectInvariantAlerts 在每日结算后检查异常：duplicate deal、active opportunity after case closed、stage out of range、negative windowDays。')

add_h2('5.4 设计预留：管道已造好，阀门未拧开')
add_table(
    ['预留项', '数据来源', '当前状态', '接入后影响'],
    [
        ['goalContext → 评分权重', 'scenario.goalContext', '只用于文案', '场景教学意图落地到结算'],
        ['boardPressureProfile → 聚焦会', 'scenario.boardPressureProfile', '完全未消费', '场景压力方向影响聚焦奖励'],
        ['ownerArchetype → satisfaction 阈值', 'ownerArchetype.trustDecayMultiplier', '全局硬编码', '不同业主判定更公平'],
        ['wordOfMouth → scoring', 'auxiliaryStats.wordOfMouth', '只在 UI 展示', '口碑建设有评分回报'],
        ['D1 + 客户质量', 'customerRuntimeState.advisorTrust/churnRisk', 'D1 只看数量', 'D1 不只看有多少人还看人多靠谱'],
    ]
)
add_para('这些不是 bug，而是迭代节奏的选择——先把管道造好，等数据积累后再决定是否拧开阀门。')

add_h2('5.5 设计哲学总结')
add_para('确定性计算 + 概率模糊 + clamp 护栏：', bold_prefix='三层架构——')
add_bullet('确定性层：所有加减乘除，产出精确的数值变化')
add_bullet('概率层：randomInt / chance / deterministicOutcomeFraction，把确定性结果模糊化')
add_bullet('护栏层：clamp / threshold / cooldown，防止极端结果')
add_para('这三层叠加产生的效果是：同样的初始条件不会产生完全相同的结局，但也不会偏离太远。')
add_para('所有引擎都是 (state) => void，就地修改 GameState，没有返回值，没有副作用。这让调试极简。代价是需要严格维护不变量。', bold_prefix='就地修改的纯函数设计——')
add_para('每天只算 2 次评分，但中间引擎可以读"昨天的 competitiveness"，避免中间态不稳定。', bold_prefix='24 步流水线 + 2 处 derivedState 刷新——')
add_para('6 套房、8 个客户、4 套竞品房源、20 个机会。每天 24 步流水线的总计算量 < 5ms。这不是"优化得好"，而是数据规模天然小——二手房经营就是这么一个"几十"的生意。', bold_prefix='数据规模是"几十"而非"几万"——')

add_highlight('终章：这个系统用最小的复杂度还原了二手房经营最核心的那个问题——在有限的信任和时间窗口内，把对的人带到桌前来。20+ 个引擎、24 步流水线、120+ 个平衡常量，最终都指向同一个目标：让玩家感受到"如果我今天不碰这套房，明天可能就来不及了"。这种紧迫感不是来自计时器，而是来自几十个数值每天在同步衰减的事实。')

# ============================================================
# 保存
# ============================================================
output_path = '/Users/jiaqi/Documents/开放日测算/王牌顾问-系统设计鉴赏.docx'
doc.save(output_path)
print(f'Word document saved to: {output_path}')
